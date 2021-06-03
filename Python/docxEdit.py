import docx

d = docx.Document("demo.docx")
print(d.paragraphs[1].text)
p = d.paragraphs[1]
p.runs[3].underline = True #Adding your edit
p.runs[3].text ='italic and underline'
d.save("outputdemo.docx")
